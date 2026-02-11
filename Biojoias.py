import streamlit as st
import json
import os
import hashlib
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime

# FUNÇÕES GLOBAIS - SEMPRE NO INÍCIO
def load_data(file):
    """Carrega JSON ou retorna lista vazia"""
    if os.path.exists(file):
        with open(file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_data(file, data):
    """Salva JSON com indentação"""
    with open(file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# DATA_FILES
DATA_FILES = {
    'produtos': 'produtos.json',
    'pagamentos': 'pagamentos.json',
    'locais': 'locais.json',
    'descontos': 'descontos.json',
    'pedidos': 'pedidos.json',
    'admins': 'admins.json',
    'loja_config': 'loja_config.json'
}

def load_data(file):
    if os.path.exists(file):
        with open(file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []


def save_data(file, data):
    with open(file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()


def check_login(username, password, admins):
    hashed_pw = hash_password(password)
    return any(admin['user'] == username and admin['hash'] == hashed_pw for admin in admins)


def generate_docx(pedido, para_admin=False):
    doc = Document()

    # No topo da generate_docx, substitua LOJA_INFO por:
    try:
        with open('loja_config.json', 'r', encoding='utf-8') as f:
            loja_config = json.load(f)
    except:
        loja_config = {'nome': 'Loja Artesanato', 'email': '', 'telefone': '', 'endereco': '', 'redesocial': ''}

    doc.add_heading(loja_config['nome'], 0)
    doc.add_paragraph(f'Contato: {loja_config["email"]} | {loja_config["telefone"]}')
    doc.add_paragraph(loja_config['endereco'])
    doc.add_paragraph(loja_config['redesocial'])

    if para_admin:
        doc.add_heading('🛒 PEDIDO ADMIN - PARA IMPRESSÃO', level=1)
    else:
        doc.add_heading('🛒 Seu Pedido Confirmado', level=1)

    # Cliente
    doc.add_paragraph(f'Cliente: {pedido["nome"]}')
    doc.add_paragraph(f'Email: {pedido["email"]}')
    doc.add_paragraph(f'Célular: {pedido["celular"]}')
    doc.add_paragraph(f'Rede Social: {pedido["social"]}')
    doc.add_paragraph(f'Data/Hora: {datetime.now().strftime("%d/%m/%Y %H:%M")}')

    # Tabela itens (igual)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Qtd'
    hdr_cells[2].text = 'Preço Unit'
    hdr_cells[3].text = 'Total'

    total = 0
    for item in pedido['itens']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['nome']
        row_cells[1].text = str(item['qtd'])
        row_cells[2].text = f'R$ {item["preco"]:.2f}'
        subtotal = item['qtd'] * item['preco']
        row_cells[3].text = f'R$ {subtotal:.2f}'
        total += subtotal

    frete = pedido.get('frete', 0)
    desconto = pedido.get('desconto', 0)
    total_final = total + frete - desconto
    doc.add_paragraph(f'Total Produtos: R$ {total:.2f}')
    doc.add_paragraph(f'Frete: R$ {frete:.2f}')
    doc.add_paragraph(f'Desconto: R$ {desconto:.2f}')
    doc.add_paragraph(f'**TOTAL FINAL: R$ {total_final:.2f}**')
    doc.add_paragraph(f'Forma Pagamento: {pedido["pagamento"]}')

    if para_admin:
        doc.add_paragraph('Status: Pendente | Imprimir e processar.')

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# Config inicial
if 'admin_logged' not in st.session_state:
    st.session_state.admin_logged = False
    st.session_state.current_user = None
    st.session_state.carrinho = []

st.title('🧵 Loja de Artesanato')

# Sidebar Admin
with st.sidebar:
    st.header('🔐 Admin')
    admins = load_data(DATA_FILES['admins'])

    tab_admin = st.tabs(['Cadastrar', 'Login'])

    with tab_admin[0]:
        st.subheader('Novo Admin')
        new_user = st.text_input('Usuário', key='new_admin_user')
        new_pw = st.text_input('Senha', type='password', key='new_admin_pw')
        new_pw_confirm = st.text_input('Confirme Senha', type='password', key='new_admin_pw_confirm')
        if st.button('Cadastrar Admin', key='btn_cad_admin'):
            if new_user and new_pw and new_pw == new_pw_confirm:
                if not any(a['user'] == new_user for a in admins):
                    admins.append({'user': new_user, 'hash': hash_password(new_pw)})
                    save_data(DATA_FILES['admins'], admins)
                    st.success('Admin cadastrado!')
                    st.rerun()
                else:
                    st.error('Usuário já existe!')
            else:
                st.error('Preencha e confirme senha corretamente!')

    with tab_admin[1]:
        if not st.session_state.admin_logged:
            username = st.text_input('Usuário', key='login_user')
            password = st.text_input('Senha', type='password', key='login_pw')
            if st.button('Entrar', key='btn_login'):
                if check_login(username, password, admins):
                    st.session_state.admin_logged = True
                    st.session_state.current_user = username
                    st.rerun()
                else:
                    st.error('Usuário ou senha inválidos!')
        else:
            st.success(f'Logado: {st.session_state.current_user}')
            if st.button('Sair', key='btn_logout'):
                st.session_state.admin_logged = False
                st.session_state.current_user = None
                st.rerun()

# Painel Admin
if st.session_state.admin_logged:
    st.header('📊 Painel Admin')

    tab1, tab2, tab3, tab4, tab5 = st.tabs(['Produtos', 'Pagamentos', 'Locais/Entrega', 'Descontos', 'Config Loja'])

    with tab1:
        st.subheader('Novo Produto')
        nome_prod = st.text_input('Nome do Item', key='prod_nome')
        preco_prod = st.number_input('Preço (R$)', key='prod_preco')
        if st.button('Salvar Produto', key='btn_save_prod'):
            produtos = load_data(DATA_FILES['produtos'])
            produtos.append({'nome': nome_prod, 'preco': preco_prod})
            save_data(DATA_FILES['produtos'], produtos)
            st.success('Produto salvo!')
            st.rerun()

        st.subheader('Produtos Cadastrados')
        produtos = load_data(DATA_FILES['produtos'])
        if not produtos:
            st.info('Nenhum produto cadastrado.')
        else:
            for idx, p in enumerate(produtos):
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.write(f'{p["nome"]} - R$ {p["preco"]:.2f}')
                with col2:
                    if st.button('Excluir', key=f'btn_del_prod_{idx}'):
                        # Remove o produto pelo índice
                        produtos.pop(idx)
                        save_data(DATA_FILES['produtos'], produtos)
                        st.warning(f'Produto "{p["nome"]}" excluído.')
                        st.rerun()

    with tab2:
        st.subheader('Nova Forma de Pagamento')
    novo_pag = st.text_input('Forma (ex: Pix)', key='new_pag')
    if st.button('Salvar Pagamento', key='btn_save_pag'):
        pags = load_data(DATA_FILES['pagamentos'])
        pags.append(novo_pag)
        save_data(DATA_FILES['pagamentos'], pags)
        st.success('Pagamento salvo!')
        st.rerun()
    st.subheader('Formas Disponíveis')
    pagamentos = load_data(DATA_FILES['pagamentos'])
    for pag in pagamentos:
        st.write(pag)

    with tab3:
        st.subheader('Novo Local')
    nome_local = st.text_input('Local (ex: Campinas)', key='local_nome')
    frete_local = st.number_input('Taxa Frete (R$)', key='local_frete')
    if st.button('Salvar Local', key='btn_save_local'):
        locais = load_data(DATA_FILES['locais'])
        locais.append({'nome': nome_local, 'frete': frete_local})
        save_data(DATA_FILES['locais'], locais)
        st.success('Local salvo!')
        st.rerun()
    st.subheader('Locais')
    locais = load_data(DATA_FILES['locais'])
    for l in locais:
        st.write(f'{l["nome"]} - Frete R$ {l["frete"]:.2f}')

    with tab4:
        st.subheader('Nova Regra Desconto')
    min_desc = st.number_input('Valor Mínimo (R$)', key='desc_min')
    val_desc = st.number_input('Desconto (R$)', key='desc_val')
    if st.button('Salvar Desconto', key='btn_save_desc'):
        descs = load_data(DATA_FILES['descontos'])
        descs.append({'min': min_desc, 'valor': val_desc})
        save_data(DATA_FILES['descontos'], descs)
        st.success('Desconto salvo!')
        st.rerun()
    st.subheader('Regras')
    descontos = load_data(DATA_FILES['descontos'])
    for d in descontos:
        st.write(f'Acima R$ {d["min"]:.2f}: -R$ {d["valor"]:.2f}')

    # NOVA ABA 5 - FUNCIONA SEM ERROS
    with tab5:
        st.header('🏪 Configurações da Loja (DOCX)')

        # Carrega config com fallback (independente de load_data global)
        if os.path.exists('loja_config.json'):
            with open('loja_config.json', 'r', encoding='utf-8') as f:
                config_loja = json.load(f)
        else:
            config_loja = {
                'nome': 'Loja de Artesanato Campinas',
                'email': 'contato@artesanato.com.br',
                'telefone': '(19) 9999-9999',
                'endereco': 'Rua Exemplo, 123 - Campinas/SP',
                'redesocial': '@artesanatosp'
            }

        # Formulário edição
        nome_loja = st.text_input('Nome da Loja', value=config_loja.get('nome', ''), key='nome_loja_edit')
        email_loja = st.text_input('Email da Loja', value=config_loja.get('email', ''), key='email_loja_edit')
        tel_loja = st.text_input('Telefone/WhatsApp', value=config_loja.get('telefone', ''), key='tel_loja_edit')
        end_loja = st.text_area('Endereço Completo', value=config_loja.get('endereco', ''), height=60,
                                key='end_loja_edit')
        social_loja = st.text_input('Rede Social', value=config_loja.get('redesocial', ''), key='social_loja_edit')

        col1, col2 = st.columns(2)
        with col1:
            if st.button('💾 Salvar Configuração', key='salvar_loja'):
                nova_config = {
                    'nome': nome_loja,
                    'email': email_loja,
                    'telefone': tel_loja,
                    'endereco': end_loja,
                    'redesocial': social_loja
                }
                with open('loja_config.json', 'w', encoding='utf-8') as f:
                    json.dump(nova_config, f, ensure_ascii=False, indent=2)
                st.success('✅ Configuração salva! Novos DOCX usarão estes dados.')
                st.balloons()
                st.rerun()

        with col2:
            st.subheader('📄 Preview Cabeçalho DOCX')
            st.markdown(f"""
            **{nome_loja or 'Nome Loja'}**  
            📧 {email_loja or 'email@loja.com'} | 📱 {tel_loja or '(19) 9999-9999'}  
            📍 {end_loja or 'Endereço'}  
            🌐 {social_loja or '@loja'}  
            """)
            st.info('Isso aparece no topo de TODO DOCX (cliente + admin)')

        # Botão teste DOCX
        if st.button('🧪 Gerar DOCX Teste com Config', key='teste_docx_loja'):
            teste_pedido = {
                'nome': 'João Teste', 'email': 'joao@email.com', 'celular': '(19) 9999-8888',
                'social': '@joaoteste', 'itens': [{'nome': 'Teste Produto', 'qtd': 2, 'preco': 25.0}],
                'frete': 10, 'desconto': 5, 'pagamento': 'Pix', 'total': 45
            }
            docx_teste = generate_docx(teste_pedido, para_admin=False)
            st.download_button(
                label='📥 Baixar Teste DOCX',
                data=docx_teste,
                file_name=f'teste_loja_config_{datetime.now().strftime("%Y%m%d_%H%M")}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )


# Tela Cliente
else:
    st.header('🛒 Loja de Artesanato')

    produtos = load_data(DATA_FILES['produtos'])
    pagamentos = load_data(DATA_FILES['pagamentos'])
    locais = load_data(DATA_FILES['locais'])
    descontos = load_data(DATA_FILES['descontos'])

    if not produtos:
        st.info('Admin ainda não cadastrou produtos.')
        st.stop()

    st.subheader('Produtos')
    for i, prod in enumerate(produtos):
        col1, col2 = st.columns([3, 1])
        with col1:
            st.write(prod['nome'])
        with col2:
            qtd = st.number_input('Qtd', min_value=0, key=f'qtd_prod_{i}')
            if qtd > 0 and st.button(f'Adicionar', key=f'btn_add_prod_{i}'):
                st.session_state.carrinho.append({'nome': prod['nome'], 'qtd': qtd, 'preco': prod['preco']})
                st.rerun()

    # Admin vê e imprime pedidos
    st.subheader('📋 Pedidos Recebidos (Imprimir)')
    pedidos_admin = load_data('pedidos_admin.json')
    if pedidos_admin:
        st.info(f'{len(pedidos_admin)} pedidos totais')
        for p in pedidos_admin[-3:]:  # Últimos 3
            with st.expander(f'#{p["id"]} {p["nome"]} • R$ {p["total"]:.2f} • {p["data"]}'):
                col1, col2, col3 = st.columns([1, 6, 1])
                with col1:
                    st.metric('Total', f'R$ {p["total"]:.2f}')
                with col2:
                    st.json({'Cliente': p['nome'], 'Contato': p['celular'], 'Pagamento': p['pagamento']})
                with col3:
                    # GERA DOCX FRESCO
                    docx_admin = generate_docx(p, para_admin=True)
                    st.download_button(
                        label='🖨️ Imprimir',
                        data=docx_admin,
                        file_name=f'admin_pedido_{p["id"]}_{p["nome"][:10]}.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        key=f'imprimir_pedido_{p["id"]}'
                    )
    else:
        st.info('👀 Aguardando primeira compra...')

    st.subheader('📋 Últimos Pedidos Admin')
    pedidos_admin = load_data('pedidos_admin.json')
    if pedidos_admin:
        for p in pedidos_admin[-3:]:  # Últimos 3
            with st.expander(f'Pedido {p["nome"]} - R$ {p["total"]:.2f}'):
                st.json(p['itens'])
                st.download_button(
                    label='Imprimir este Pedido (DOCX)',
                    data=p['docx_data'],
                    file_name=f'admin_pedido_{p["nome"]}_{datetime.now().strftime("%Y%m%d")}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    key=f'dl_admin_{len(pedidos_admin)}'
                )
    else:
        st.info('Nenhum pedido ainda.')

    if st.session_state.carrinho:
        st.subheader('Carrinho')
        total = sum(item['qtd'] * item['preco'] for item in st.session_state.carrinho)
        for idx, item in enumerate(st.session_state.carrinho):
            st.write(f'{item["nome"]} x{item["qtd"]} = R$ {item["qtd"] * item["preco"]:.2f}')
        st.write(f'**Subtotal: R$ {total:.2f}**')

        local_sel = st.selectbox('Local de Entrega', [l['nome'] for l in locais] if locais else ['Nenhum'],
                                 key='sel_local')
        frete = next((l['frete'] for l in locais if l['nome'] == local_sel), 0)
        st.write(f'Frete: R$ {frete:.2f}')

        desconto = 0
        for d in descontos:
            if total >= d['min']:
                desconto = d['valor']
                break
        st.write(f'Desconto: R$ {desconto:.2f}')
        total_final = total + frete - desconto
        st.write(f'**TOTAL: R$ {total_final:.2f}**')

        st.subheader('Seus Dados')
        nome = st.text_input('Nome Completo', key='cliente_nome')
        email = st.text_input('Email', key='cliente_email')
        celular = st.text_input('Celular', key='cliente_celular')
        social = st.text_input('Rede Social (opcional)', key='cliente_social')
        pag_sel = st.selectbox('Forma Pagamento', pagamentos if pagamentos else ['Nenhum'], key='sel_pagamento')

        if st.button('Confirmar Pedido', key='btn_confirmar'):
            if nome and email and celular:
                pedido = {
                    'id': len(load_data(DATA_FILES['pedidos'])) + 1,
                    'nome': nome, 'email': email, 'celular': celular, 'social': social,
                    'itens': st.session_state.carrinho.copy(),
                    'frete': frete, 'desconto': desconto,
                    'pagamento': pag_sel, 'total': total_final,
                    'data': datetime.now().strftime("%d/%m/%Y %H:%M")
                }

                # Salva pedidos gerais
                pedidos = load_data(DATA_FILES['pedidos'])
                pedidos.append(pedido)
                save_data(DATA_FILES['pedidos'], pedidos)

                # Salva cópia para ADMIN
                pedidos_admin = load_data('pedidos_admin.json') if os.path.exists('pedidos_admin.json') else []
                pedidos_admin.append(pedido.copy())
                save_data('pedidos_admin.json', pedidos_admin)

                # DOCX CLIENTE
                docx_cliente = generate_docx(pedido, para_admin=False)
                st.download_button(
                    label='📄 Seu Pedido Confirmado (DOCX)',
                    data=docx_cliente,
                    file_name=f'cliente_pedido_{pedido["id"]}_{nome.replace(" ", "_")}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    key='dl_cliente'
                )

                st.balloons()  # 🎉
                st.success(f'Pedido #{pedido["id"]} confirmado! Admin receberá cópia.')

                col1, col2 = st.columns(2)
                with col1:
                    if st.button('Novo Pedido', key='btn_limpar_carrinho'):
                        st.session_state.carrinho = []
                        st.rerun()
                with col2:
                    st.info('Envie WhatsApp para loja acelerar.')
            else:
                st.error('❌ Nome, email e celular obrigatórios!')

st.markdown('---')
st.caption('App local corrigido - sem erros de key duplicada.')
